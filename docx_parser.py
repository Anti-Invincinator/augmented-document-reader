#!/usr/bin/env python3
"""
Financial Document Named Entity Extractor
Extracts key financial terms and entities from structured financial documents (docx format)
"""

import re
from datetime import datetime
from typing import Dict, Any, Optional, List
import json
from docx import Document
from pathlib import Path


class FinancialDocumentParser:
    """
    Rule-based parser for extracting named entities from financial derivative documents
    """
    
    def __init__(self):
        # Define patterns for various financial entities
        self.patterns = {
            'trade_id': r'([A-Z]{2}\d{4})',
            'isin': r'([A-Z]{2}\d{9}[A-Z0-9])',
            'reuters_code': r'([A-Z]+\.[A-Z]+)',
            'currency': r'\b(EUR|USD|GBP|JPY|CHF|CAD|AUD)\b',
            'percentage': r'(\d+(?:\.\d+)?%)',
            'amount_with_currency': r'((?:EUR|USD|GBP|JPY|CHF|CAD|AUD)\s+[\d,]+(?:\.\d+)?(?:\s+(?:million|billion|thousand))?)',
            'date': r'(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})',
            'time': r'(\d{2}:\d{2}:\d{2})',
            'exchange_code': r'\b(XETRA|NYSE|NASDAQ|LSE|TSE)\b',
            'business_day_convention': r'\b(TARGET|FOLLOWING|PRECEDING|MODIFIED FOLLOWING)\b'
        }
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract all text content from docx file"""
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error reading docx file: {e}")
            return ""
    
    def extract_field_value_pairs(self, text: str) -> Dict[str, str]:
        """Extract field-value pairs from structured text"""
        pairs = {}
        
        # Common field patterns in financial documents
        field_patterns = {
            'party_a': r'Party A[:\s]+(.+?)(?:\n|$)',
            'party_b': r'Party B[:\s]+(.+?)(?:\n|$)',
            'trade_date': r'Trade Date[:\s]+(.+?)(?:\n|$)',
            'trade_time': r'Trade Time[:\s]+(.+?)(?:\n|$)',
            'initial_valuation_date': r'Initial Valuation Date[:\s]+(.+?)(?:\n|$)',
            'effective_date': r'Effective Date[:\s]+(.+?)(?:\n|$)',
            'notional_amount': r'Notional Amount[:\s\(N\)]*(.+?)(?:\n|$)',
            'upfront_payment': r'Upfront Payment[:\s]+(.+?)(?:\n|$)',
            'valuation_date': r'Valuation Date[:\s]+(.+?)(?:\n|$)',
            'termination_date': r'Termination Date[:\s]+(.+?)(?:\n|$)',
            'underlying': r'Underlying[:\s]+(.+?)(?:\n|$)',
            'exchange': r'Exchange[:\s]+(.+?)(?:\n|$)',
            'coupon': r'Coupon[:\s\(C\)]*(.+?)(?:\n|$)',
            'barrier': r'Barrier[:\s\(B\)]*(.+?)(?:\n|$)',
            'calculation_agent': r'Calculation Agent[:\s]+(.+?)(?:\n|$)',
            'isda_documentation': r'ISDA Documentation[:\s]+(.+?)(?:\n|$)',
            'business_day': r'Business Day[:\s]+(.+?)(?:\n|$)'
        }
        
        for field, pattern in field_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                value = match.group(1).strip()
                # Clean up common formatting artifacts
                value = re.sub(r'\*+', '', value)  # Remove asterisks
                value = re.sub(r'TBD', 'To Be Determined', value)
                pairs[field] = value
        
        return pairs
    
    def extract_named_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract named entities using regex patterns"""
        entities = {}
        
        for entity_type, pattern in self.patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                # Remove duplicates while preserving order
                unique_matches = list(dict.fromkeys(matches))
                entities[entity_type] = unique_matches
        
        return entities
    
    def extract_financial_terms(self, text: str) -> Dict[str, Any]:
        """Extract specific financial terms and calculate derived values"""
        terms = {}
        
        # Extract barrier percentage
        barrier_match = re.search(r'(\d+(?:\.\d+)?)\s*%.*?of.*?Share', text, re.IGNORECASE)
        if barrier_match:
            terms['barrier_percentage'] = float(barrier_match.group(1))
        
        # Extract notional amount
        notional_match = re.search(r'(EUR|USD|GBP)\s+([\d,]+(?:\.\d+)?)\s*(million|billion|thousand)?', text, re.IGNORECASE)
        if notional_match:
            currency = notional_match.group(1)
            amount = float(notional_match.group(2).replace(',', ''))
            multiplier = notional_match.group(3)
            
            if multiplier:
                multiplier = multiplier.lower()
                if multiplier == 'million':
                    amount *= 1_000_000
                elif multiplier == 'billion':
                    amount *= 1_000_000_000
                elif multiplier == 'thousand':
                    amount *= 1_000
            
            terms['notional_amount'] = {
                'currency': currency,
                'amount': amount,
                'formatted': f"{currency} {amount:,.2f}"
            }
        
        return terms
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """Main parsing method that combines all extraction techniques"""
        
        # Extract text from document
        text = self.extract_text_from_docx(file_path)
        
        if not text:
            return {"error": "Could not extract text from document"}
        
        # Extract structured data
        field_values = self.extract_field_value_pairs(text)
        named_entities = self.extract_named_entities(text)
        financial_terms = self.extract_financial_terms(text)
        
        # Create comprehensive result
        result = {
            "document_info": {
                "file_path": file_path,
                "extraction_timestamp": datetime.now().isoformat(),
                "text_length": len(text)
            },
            "structured_fields": field_values,
            "named_entities": named_entities,
            "financial_terms": financial_terms,
            "raw_text": text  # Include for debugging/verification
        }
        
        return result
    
    def save_results(self, results: Dict[str, Any], output_path: str, format: str = 'json'):
        """Save extraction results to file"""
        
        if format.lower() == 'json':
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
        
        elif format.lower() == 'txt':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("FINANCIAL DOCUMENT EXTRACTION RESULTS\n")
                f.write("=" * 50 + "\n\n")
                
                # Write structured fields
                f.write("STRUCTURED FIELDS:\n")
                f.write("-" * 20 + "\n")
                for key, value in results.get('structured_fields', {}).items():
                    f.write(f"{key.replace('_', ' ').title()}: {value}\n")
                
                f.write("\nNAMED ENTITIES:\n")
                f.write("-" * 15 + "\n")
                for entity_type, entities in results.get('named_entities', {}).items():
                    f.write(f"{entity_type.replace('_', ' ').title()}: {', '.join(entities)}\n")
                
                f.write("\nFINANCIAL TERMS:\n")
                f.write("-" * 16 + "\n")
                for key, value in results.get('financial_terms', {}).items():
                    f.write(f"{key.replace('_', ' ').title()}: {value}\n")


def main():
    """Example usage of the parser"""
    
    # Initialize parser
    parser = FinancialDocumentParser()
    
    # Example usage - replace with actual file path
    input_file = "ZF4894_ALV_07Aug2026_physical.docx"
    
    if Path(input_file).exists():
        # Parse the document
        results = parser.parse_document(input_file)
        
        # Save results in multiple formats
        base_name = Path(input_file).stem
        parser.save_results(results, f"{base_name}_extracted.json", 'json')
        parser.save_results(results, f"{base_name}_extracted.txt", 'txt')
        
        print("Extraction completed successfully!")
        print(f"Results saved to {base_name}_extracted.json and {base_name}_extracted.txt")
        
        # Print summary
        print("\nExtraction Summary:")
        print(f"Structured fields found: {len(results.get('structured_fields', {}))}")
        print(f"Named entity types found: {len(results.get('named_entities', {}))}")
        print(f"Financial terms extracted: {len(results.get('financial_terms', {}))}")
        
    else:
        print(f"File {input_file} not found. Please provide the correct path.")
        
        # Show example of how to use with different file
        print("\nTo use with your file:")
        print("parser = FinancialDocumentParser()")
        print("results = parser.parse_document('your_file.docx')")
        print("parser.save_results(results, 'output.json')")


if __name__ == "__main__":
    main()